from django.shortcuts import render, redirect
from django.core.exceptions import ValidationError
from django.contrib import messages
from .models import Resume
import re
from pdfminer.high_level import extract_text
import docx
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
import json
import os
import google.generativeai as genai
from django.conf import settings
from django.core.files.uploadedfile import UploadedFile
import logging
import tempfile
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

# Set up logging
logger = logging.getLogger(__name__)

# Configure Gemini API
try:
    genai.configure(api_key=settings.GOOGLE_GEMINI_API_KEY)
except AttributeError:
    logger.error("Google Gemini API key is missing in settings.")
    raise Exception("Google Gemini API key not configured.")

# Cache file for skills and degrees
CACHE_FILE = os.path.join(settings.BASE_DIR, "cached_skills_degrees.json")
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB


def get_cached_data():
    """Load skills and degrees from cache or fetch from Gemini API if not available."""
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r") as file:
                logger.debug(f"Loading cached data from {CACHE_FILE}")
                data = json.load(file)
                logger.debug(f"Loaded data: {data}")
                return data
        except (json.JSONDecodeError, IOError) as e:
            logger.error(f"Failed to load cached data: {e}")
            return fetch_skills_degrees_from_gemini()
    logger.info(f"Cache file {CACHE_FILE} not found. Fetching from Gemini API.")
    return fetch_skills_degrees_from_gemini()

def fetch_skills_degrees_from_gemini():
    """Fetch skills and degrees from Gemini API and cache them."""
    prompt = """
    Generate two JSON lists and return the result as a valid JSON string only, with no additional text, markdown, or formatting:
    - "skills": A list of 5000 job-related skills across all industries, including both soft skills (e.g., Communication, Leadership) and technical skills (e.g., Python, Java, HTML, CSS, JavaScript, SQL, Django, React, Git, Github, MySQL, MongoDB, Tailwind CSS, Swift, Go, R, Node.js, AWS, Docker, Kubernetes, Machine Learning, Data Science).
    - "degrees": A list of 150 academic degrees worldwide, including variations of engineering degrees (e.g., Bachelor of Engineering, BEng, Master of Engineering, MEng, Computer Engineering) and other common degrees (e.g., B.Tech, M.Sc, PhD, B.A., M.A.).
    """
    try:
        logger.debug("Attempting to fetch data from Gemini API...")
        model = genai.GenerativeModel("gemini-1.5-pro-latest")
        response = model.generate_content(prompt)
        logger.debug(f"Raw response from Gemini API: {response.text}")
        # Clean the response to extract JSON
        cleaned_response = response.text.strip()
        cleaned_response = re.sub(r'```json\s*\n|\n\s*```', '', cleaned_response, flags=re.MULTILINE)
        cleaned_response = re.sub(r'```\s*\n|\n\s*```', '', cleaned_response, flags=re.MULTILINE)
        cleaned_response = re.sub(r'^.*?{', '{', cleaned_response, flags=re.DOTALL)
        cleaned_response = re.sub(r'}[^}]*$', '}', cleaned_response, flags=re.DOTALL)
        cleaned_response = cleaned_response.strip()
        logger.debug(f"Cleaned response: {cleaned_response}")
        # Parse the cleaned response as JSON
        data = json.loads(cleaned_response)
        logger.debug(f"Parsed JSON data: {data}")
        if not isinstance(data, dict) or "skills" not in data or "degrees" not in data:
            logger.error(f"Gemini API returned invalid data: {data}")
            return {"skills": [], "degrees": []}
        logger.debug(f"Writing to cache file: {CACHE_FILE}")
        with open(CACHE_FILE, "w") as file:
            json.dump(data, file, indent=4)
        logger.info(f"Successfully created cache file: {CACHE_FILE}")
        return data
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {e}, Raw response: {response.text}")
        return {"skills": [], "degrees": []}
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        return {"skills": [], "degrees": []}

CACHED_DATA = get_cached_data()
ALL_SKILLS = set(CACHED_DATA.get("skills", []))
ALL_DEGREES = set(CACHED_DATA.get("degrees", []))
logger.debug(f"ALL_SKILLS: {ALL_SKILLS}")
logger.debug(f"ALL_DEGREES: {ALL_DEGREES}")


# Path to Tesseract executable (make configurable)
TESSERACT_CMD = getattr(settings, 'TESSERACT_CMD', r'C:\Program Files\Tesseract-OCR\tesseract.exe')
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# File validation
def validate_file(uploaded_file: UploadedFile):
    """Validate file type and size."""
    if not uploaded_file:
        raise ValidationError("No file uploaded.")
    
    file_type = uploaded_file.name.split('.')[-1].lower()
    if file_type not in ALLOWED_EXTENSIONS:
        raise ValidationError(f"Unsupported file format: {file_type}. Only PDF and DOCX are allowed.")
    
    if uploaded_file.size > MAX_FILE_SIZE:
        raise ValidationError(f"File too large. Maximum size allowed is {MAX_FILE_SIZE / (1024 * 1024)} MB.")

# PDF text extraction
def extract_text_from_pdf(pdf_file: UploadedFile) -> str:
    """Extract text from a PDF file using pdfminer and OCR."""
    try:
        # First try pdfminer
        pdf_bytes = pdf_file.read()
        pdf_stream = io.BytesIO(pdf_bytes)
        text = extract_text(pdf_stream)
        
        # If pdfminer fails or returns very little text, try OCR
        if not text or len(text.strip()) < 100:  # If less than 100 characters
            logger.debug("PDF text extraction returned little text, trying OCR")
            pdf_file.seek(0)  # Reset file pointer
            return extract_text_with_ocr(pdf_file)
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""

# OCR for scanned PDFs
def extract_text_with_ocr(pdf_file: UploadedFile) -> str:
    """Extract text from PDF using OCR with improved settings."""
    try:
        pdf_bytes = pdf_file.read()
        images = convert_from_bytes(pdf_bytes)
        text = ""
        
        for img in images:
            # Convert to grayscale for better OCR
            img = img.convert('L')
            # Enhance contrast
            img = img.point(lambda x: 0 if x < 128 else 255, '1')
            
            # Try different OCR configurations
            configs = [
                '--psm 6',  # Assume a single uniform block of text
                '--psm 3',  # Fully automatic page segmentation
                '--psm 4',  # Assume a single column of text
            ]
            
            for config in configs:
                try:
                    ocr_text = pytesseract.image_to_string(img, config=config)
                    if ocr_text.strip():
                        text += ocr_text + "\n"
                        break
                except Exception as e:
                    logger.error(f"OCR failed with config {config}: {e}")
                    continue
        
        return text.strip()
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return ""

# DOCX text extraction
def extract_text_from_docx(docx_file: UploadedFile) -> str:
    """Extract text from a DOCX file with improved handling."""
    try:
        doc = docx.Document(docx_file)
        text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ""

# Clean and preprocess text
def preprocess_text(text: str) -> str:
    """Clean and preprocess extracted text with improved handling."""
    if not text:
        return ""
    
    # Replace multiple spaces and newlines with single space
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep email and phone number patterns
    text = re.sub(r'[^a-zA-Z0-9@.,+()&:;\s-]', ' ', text)
    
    # Normalize spaces around special characters
    text = re.sub(r'\s*([@.,+()&:;-])\s*', r'\1', text)
    
    return text.strip()

# Field extraction
def extract_skills(text: str) -> str:
    """Extract skills by matching against the cached skills list."""
    if not ALL_SKILLS:
        logger.warning("Skills list is empty. Check Gemini API or cache.")
        return "N/A"
    text_lower = text.lower()
    matched_skills = set()
    for skill in ALL_SKILLS:
        skill_lower = skill.lower()
        # Match whole words to avoid partial matches (e.g., "Java" in "JavaScript")
        # Use word boundaries and ensure the skill is a standalone word or phrase
        pattern = r'\b' + re.escape(skill_lower) + r'\b'
        if re.search(pattern, text_lower):
            matched_skills.add(skill)
    logger.debug(f"Matched Skills: {matched_skills}")
    return ', '.join(sorted(matched_skills)) if matched_skills else "N/A"

def extract_education(text: str) -> str:
    """Extract education by matching against the cached degrees list."""
    if not ALL_DEGREES:
        logger.warning("Degrees list is empty. Check Gemini API or cache.")
        return "N/A"
    text_lower = text.lower()
    matched_degrees = set()
    for degree in ALL_DEGREES:
        degree_lower = degree.lower()
        # Match whole phrases with word boundaries
        pattern = r'\b' + re.escape(degree_lower) + r'\b'
        if re.search(pattern, text_lower):
            matched_degrees.add(degree)
        # Handle variations for engineering degrees
        elif "bachelor of engineering" in degree_lower and "bachelor of engineering" in text_lower:
            matched_degrees.add(degree)
        elif "beng" in degree_lower and "bachelor of engineering" in text_lower:
            matched_degrees.add(degree)
        # Match "Computer Engineering" specifically
        elif "computer engineering" in degree_lower and "computer engineering" in text_lower:
            matched_degrees.add(degree)
    logger.debug(f"Matched Degrees: {matched_degrees}")
    return ', '.join(sorted(matched_degrees)) if matched_degrees else "N/A"

def extract_name(text: str) -> str:
    """Extract name from resume text using common resume patterns."""
    try:
        # Common patterns found in resumes
        patterns = [
            r'Name[:\s]+([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)',  # Name: John Doe
            r'^([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)',           # John Doe at start
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)\s+Resume',   # John Doe Resume
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)\s+[A-Z][a-z]+',  # John Doe followed by another word
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)\s*\|',       # John Doe | 
            r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)\s*[-–]',     # John Doe - 
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Basic validation - name should be at least 3 chars and contain at least one space
                if len(name) >= 3 and ' ' in name:
                    return name
    except Exception as e:
        logger.error(f"Error extracting name: {e}")
    return "N/A"

def extract_email(text: str) -> str:
    """Extract email from text using common email patterns found in resumes."""
    try:
        # Focus on the top portion of the resume (first 20%)
        top_portion = text[:len(text)//5]
        logger.debug(f"Searching for email in top portion: {top_portion[:200]}...")
        
        # Common email patterns in resumes
        patterns = [
            r'Email[:\s]*([\w\.-]+@[\w\.-]+\.\w+)',  # After "Email:" or similar
            r'[\w\.-]+@[\w\.-]+\.(?:com|org|net|edu|gov)',  # Common TLDs
            r'[\w\.-]+@[\w\.-]+\.(?:co\.uk|co\.in|ac\.uk)',  # Country-specific TLDs
            r'[\w\.-]+@[\w\.-]+\.(?:io|ai|me)',  # Modern TLDs
            r'[\w\.-]+@[\w\.-]+\.(?:co|in|us|ca|au)',  # More country TLDs
            r'[\w\.-]+@[\w\.-]+\.(?:biz|info|name)',  # Other common TLDs
            r'[\w\.-]+@[\w\.-]+\.(?:tech|dev|app)',  # Tech-related TLDs
        ]
        
        # First try to find email in the first few lines
        lines = top_portion.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if not line:
                continue
                
            # Skip lines that are too short or don't look like contact info
            if len(line) < 5 or line.isupper() or line.startswith(('SUMMARY', 'OBJECTIVE', 'EXPERIENCE', 'EDUCATION')):
                continue
                
            for pattern in patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    email = match.group(1) if match.groups() else match.group(0)
                    # Basic email validation
                    if '@' in email and '.' in email and len(email) >= 5:
                        logger.debug(f"Found email: {email}")
                        return email.lower().strip()
        
        logger.debug("No email pattern matched in top portion")
    except Exception as e:
        logger.error(f"Error extracting email: {e}")
    return "N/A"

def extract_phone(text: str) -> str:
    """Extract phone number from text, handling various formats including country codes."""
    try:
        # Focus on the top portion of the resume (first 20%)
        top_portion = text[:len(text)//5]
        logger.debug(f"Searching for phone in top portion: {top_portion[:200]}...")
        
        # Common phone patterns in resumes
        patterns = [
            r'Phone[:\s]*([+\d\s\-()]+)',  # After "Phone:" or similar
            r'Tel[:\s]*([+\d\s\-()]+)',    # After "Tel:" or similar
            r'Contact[:\s]*([+\d\s\-()]+)', # After "Contact:" or similar
            r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # Standard phone with optional country code
            r'(\+\d{1,3}[-.\s]?)?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',        # Standard phone without parentheses
            r'(\+\d{1,3}[-.\s]?)?\d{10}',                               # 10 digits with optional country code
            r'(\+\d{1,3}[-.\s]?)?\d{4}[-.\s]?\d{3}[-.\s]?\d{3}',        # Alternative format
            # New pattern for country code with space
            r'(\+\d{1,3})\s*(\d{10})',                                  # Country code with space
            r'(\+\d{1,3})\s*(\d{4}[-.\s]?\d{3}[-.\s]?\d{3})',           # Country code with space and formatted number
        ]
        
        # First try to find phone in the first few lines
        lines = top_portion.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if not line:
                continue
                
            # Skip lines that are too short or don't look like contact info
            if len(line) < 5 or line.isupper() or line.startswith(('SUMMARY', 'OBJECTIVE', 'EXPERIENCE', 'EDUCATION')):
                continue
                
            for pattern in patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    if len(match.groups()) > 1:  # For patterns with country code and number as separate groups
                        country_code = match.group(1)
                        number = match.group(2)
                        phone = f"{country_code}{number}"
                    else:
                        phone = match.group(1) if match.groups() else match.group(0)
                    
                    # Clean the phone number - keep only digits and +
                    cleaned_phone = re.sub(r'[^\d+]', '', phone)
                    
                    # Basic validation - should be at least 10 digits after country code
                    if len(cleaned_phone) >= 10:
                        logger.debug(f"Found phone: {cleaned_phone}")
                        return cleaned_phone
        
        logger.debug("No phone pattern matched in top portion")
    except Exception as e:
        logger.error(f"Error extracting phone: {e}")
    return "N/A"

# Main parsing logic
def upload_resume(request):
    """Handle resume upload, parsing, and storage."""
    if request.method != 'POST':
        try:
            return render(request, 'resume_parser/upload_resume.html')
        except Exception as e:
            logger.error(f"Error rendering upload_resume.html: {e}")
            messages.error(request, "Error loading upload page. Please try again.")
            return redirect('upload_resume')

    try:
        # Validate file
        uploaded_file = request.FILES.get('resume')
        validate_file(uploaded_file)

        file_type = uploaded_file.name.split('.')[-1].lower()
        file_text = ""

        # Extract text based on file type
        if file_type == 'pdf':
            file_text = extract_text_from_pdf(uploaded_file)
            if not file_text.strip():
                uploaded_file.seek(0)  # Reset file pointer
                file_text = extract_text_with_ocr(uploaded_file)
        elif file_type == 'docx':
            file_text = extract_text_from_docx(uploaded_file)

        if not file_text.strip():
            messages.error(request, "Unable to extract text from the file. Please try a different file.")
            return redirect('upload_resume')

        # Preprocess and parse the text
        cleaned_text = preprocess_text(file_text)
        logger.debug(f"Cleaned Text: {cleaned_text}")

        parsed_data = {
            'name': extract_name(cleaned_text),
            'email': extract_email(cleaned_text),
            'phone': extract_phone(cleaned_text),
            'skills': extract_skills(cleaned_text) or "N/A",
            'education': extract_education(cleaned_text) or "N/A",
        }

        logger.debug(f"Parsed Data: {parsed_data}")

        # Save to database with validation
        try:
            resume_instance = Resume(**parsed_data)
            resume_instance.full_clean()  # Validate the instance
            resume_instance.save()
        except ValidationError as ve:
            logger.error(f"Validation error saving resume: {ve}")
            messages.error(request, f"Invalid data in resume: {str(ve)}. Please check the file and try again.")
            return redirect('upload_resume')
        except Exception as e:
            logger.error(f"Error saving to database: {e}")
            messages.error(request, "Error saving resume data. Please try again.")
            return redirect('upload_resume')

        messages.success(request, "Resume parsed and saved successfully!")
        try:
            return render(request, 'resume_parser/resume_parsed.html', {'parsed_data': parsed_data})
        except Exception as e:
            logger.error(f"Error rendering resume_parsed.html: {e}")
            messages.error(request, "Resume saved, but there was an error displaying the results.")
            return redirect('upload_resume')

    except ValidationError as ve:
        messages.error(request, str(ve))
        return redirect('upload_resume')
    except Exception as e:
        logger.error(f"Unexpected error in upload_resume: {e}")
        messages.error(request, "An unexpected error occurred. Please try again later.")
        return redirect('upload_resume')

@csrf_exempt
@require_http_methods(["POST"])
def api_upload_resume(request):
    """API endpoint for resume upload and parsing."""
    try:
        logger.debug("Received file upload request")
        
        # Validate file
        uploaded_file = request.FILES.get('resume')
        if not uploaded_file:
            logger.error("No file received in request")
            return JsonResponse({
                'error': 'No file uploaded'
            }, status=400)

        logger.debug(f"Received file: {uploaded_file.name} ({uploaded_file.size} bytes)")
        
        validate_file(uploaded_file)

        file_type = uploaded_file.name.split('.')[-1].lower()
        file_text = ""

        # Extract text based on file type
        if file_type == 'pdf':
            logger.debug("Processing PDF file")
            file_text = extract_text_from_pdf(uploaded_file)
            if not file_text.strip():
                logger.debug("PDF text extraction failed, trying OCR")
                uploaded_file.seek(0)  # Reset file pointer
                file_text = extract_text_with_ocr(uploaded_file)
        elif file_type == 'docx':
            logger.debug("Processing DOCX file")
            file_text = extract_text_from_docx(uploaded_file)

        if not file_text.strip():
            logger.error("Failed to extract text from file")
            return JsonResponse({
                'error': 'Unable to extract text from the file. Please try a different file.'
            }, status=400)

        # Preprocess and parse the text
        cleaned_text = preprocess_text(file_text)
        logger.debug(f"Cleaned Text length: {len(cleaned_text)}")

        parsed_data = {
            'name': extract_name(cleaned_text),
            'email': extract_email(cleaned_text),
            'phone': extract_phone(cleaned_text),
            'skills': extract_skills(cleaned_text) or "N/A",
            'education': extract_education(cleaned_text) or "N/A",
        }

        logger.debug(f"Parsed Data: {parsed_data}")

        # Save to database
        resume_instance = Resume(**parsed_data)
        resume_instance.full_clean()
        resume_instance.save()

        return JsonResponse({
            'status': 'success',
            'message': 'Resume parsed successfully',
            'data': parsed_data
        })

    except ValidationError as ve:
        logger.error(f"Validation error: {str(ve)}")
        return JsonResponse({
            'error': str(ve)
        }, status=400)
    except Exception as e:
        logger.error(f"Unexpected error in api_upload_resume: {str(e)}", exc_info=True)
        return JsonResponse({
            'error': 'An unexpected error occurred. Please try again later.'
        }, status=500)


